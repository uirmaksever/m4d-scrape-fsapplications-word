import docx
import pandas as pd
import os

# This script deals with application forms which are in the form of word documents.


filenames = open("filenames.txt").readlines()
# Get rid of \n's
filenames = [line.rstrip() for line in filenames]



# Look for designated locations, extract info and create entry out of it
def get_info(file_name):
    file_path = "sadece_basvurular/" + file_name
    try:
        with open(file_path, mode="rb") as file:
            file_docx = docx.Document(file)

        tables = file_docx.tables
        info = {}
        info["dosyaadi"] = file_name
        info["basvuru_sahibi"] = tables[0].rows[1].cells[2].text
        info["adres"] = tables[0].rows[3].cells[2].text
        info["iletisim_kisisi"] = tables[1].rows[0].cells[2].text
        info["e_posta_adresi"] = tables[1].rows[4].cells[2].text
        info["telefon"] = tables[1].rows[3].cells[2].text
        info["calisma_adi"] = tables[2].rows[1].cells[2].text
        info["calisma_ozeti"] = tables[2].rows[2].cells[2].text
        for key in info.keys():
            value = info[key]
            value = value.replace("\n","")
            info[key] = value

        records.append(info)
        print(info)
    except docx.Document.PackageNotFoundError:
        empty = "EMPTY"
        info["dosyaadi"] = file_name
        info["basvuru_sahibi"] = empty
        info["adres"] = empty
        info["iletisim_kisisi"]= empty
        info["e_posta_adresi"] = empty
        info["telefon"] = empty
        info["calisma_adi"] = empty
        info["calisma_ozeti"] = empty
        records.append(info)
        print("EMPTY!!!!", file_name)


# Iterate over items and get an xlsx export
records = []
for file_name in filenames:
    get_info(file_name)
record_pd = pd.DataFrame(records)
record_pd.to_excel("export.xlsx")

####

# Total budget amount was not in a table. Iterate over all paragraphs and take it
def find_budget(doc_object):
    doc_paragraphs = doc_object.paragraphs
    found = []
    for paragraph in doc_paragraphs:
        text = paragraph.text
        search_text = "Tüm faaliyetler için talep edilen toplam bütçe"
        if text.__contains__(search_text):
            found.append(text)
            print("found", text)
        else:
            pass
    return found

# Iteration
budgets = []
for file_name in filenames:
    file_path = "sadece_basvurular/" + file_name
    budget_info = {}
    print(file_name)
    with open(file_path, mode="rb") as file:
        doc_object = docx.Document(file)

    result = find_budget(doc_object)
    result = str(result)
    # result = result[0].replace(search_text, "")
    budget_info["file_name"] = file_name
    budget_info["budget"] = result
    budgets.append(budget_info)

budgets_pd = pd.DataFrame(budgets)
budgets_pd.to_excel("budgets.xlsx")

#######

# Similar to budget, application type could only be figured out by document
# title, which is also another paragraph item.
# Look for both titles on each iteration, save it.
def detect_application_type(doc_object):
    doc_paragraphs = doc_object.paragraphs
    found = []
    for paragraph in doc_paragraphs:
        text = paragraph.text
        kurumsal = "YAPISAL GÜÇLENDİRME DESTEK PAKETİ BAŞVURU FORMU"
        bireysel = "MESLEKİ SÜRDÜRÜLEBİLİRLİK DESTEK PAKETİ BAŞVURU FORMU"
        if text.__contains__(bireysel):
            found.append(bireysel)
        elif text.__contains__(kurumsal):
            found.append(kurumsal)
        else:
            pass
    return found

# Iteration
application_types = []
for file_name in filenames:
    file_path = "sadece_basvurular/" + file_name

    deneme = docx.Document(open(file_path, mode="rb"))
    application_type = detect_application_type(deneme)
    info = {}
    info["filename"] = file_name
    info["type"] = application_type
    application_types.append(info)
    print(application_type)

application_types_pd = pd.DataFrame(application_types)
application_types_pd.to_excel("application_types.xlsx")


yapisalornek = "10-446062.docx"
bireyselornek = "12-244739.docx"

def docx_from_filename(filename):
    prefix = "sadece_basvurular/"
    file_path = prefix + filename

    with open(file_path, mode="rb") as doc_file:
        doc_object = docx.Document(doc_file)

    return doc_object

ornek = docx_from_filename(yapisalornek)
budgets = get_budget_tables(ornek, yapisalornek)

budgets_list = []
for filename in filenames:
    docx_object = docx_from_filename(filename)
    budgets_pd = get_budget_tables(docx_object, filename)
    budgets_list.append(budgets_pd)

dfs = pd.concat(budgets_list)
dfs.to_excel("all_applications_budget.xlsx")

def get_budget_tables(doc_object, filename):

    budget_tables = []
    for table in doc_object.tables:
        number_of_rows = len(table.rows)

        iteration = 0
        for row in table.rows:
            row_number_of_columns = len(row.cells)

            if row_number_of_columns == 5:
                row_dict = {}
                row_dict["filename"] = filename
                row_dict["application_type"] = detect_application_type(doc_object)
                row_dict["table_id"] = table.rows[0].cells[0].text
                row_dict["id"] = row.cells[0].text
                row_dict["gider_kalemi"] = row.cells[1].text
                row_dict["birim"] = row.cells[2].text
                row_dict["birim_maliyet"] = row.cells[3].text
                row_dict["toplam"] = row.cells[4].text
                budget_tables.append(row_dict)
        print(sample_row_number_of_columns)
        iteration = iteration + 1

    budget_tables_pd = pd.DataFrame(budget_tables)

    return budget_tables_pd


#### PDF applications
pdf_prefix = "basvurular_pdf/"
pdf_filename = "50-527544.pdf"

import camelot


pdf = camelot.read_pdf(pdf_prefix + pdf_filename)
pdf
